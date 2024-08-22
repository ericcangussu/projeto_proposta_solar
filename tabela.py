import streamlit as st
import pandas as pd
from datetime import datetime
import locale


#Marco 1
# Título da aplicação
st.title('Formulário de Proposta Solar')

# Seção: Dados Iniciais
st.subheader('Dados Iniciais')

# Pergunta se a proposta é referente ao dia de hoje
use_today_date = st.radio('A proposta é referente ao dia de hoje?', ('Sim', 'Não'))

if use_today_date == 'Sim':
    selected_date = datetime.today()
    dia = selected_date.strftime("%d")
    mes = selected_date.strftime("%B")  # Nome completo do mês
    ano = selected_date.strftime("%Y")
else:
    selected_date = st.date_input('Selecione a data:')
    dia = selected_date.strftime("%d")
    mes = selected_date.strftime("%B")  # Nome completo do mês
    ano = selected_date.strftime("%Y")

cliente = st.text_input('Cliente:')
obra = st.text_input('Obra:')
numero_or = st.text_input('Número da OR:')
revisao = st.text_input('Revisão:')
cotacao_dolar = st.number_input('Cotação do Dólar:', min_value=0.0, value=5.4, format="%.2f")

# Inicializa o st.session_state['unique_key'] se ainda não estiver presente
# Inicializa o st.session_state['unique_key'] se ainda não estiver presente
if 'unique_key' not in st.session_state:
    st.session_state['unique_key'] = 0

# Função para gerar uma chave única
def gerar_chave_unica(prefixo):
    chave_unica = f"{prefixo}_{st.session_state['unique_key']}"
    st.session_state['unique_key'] += 1
    return chave_unica

# Inicializa a lista de itens configurados se ainda não estiver presente
if 'itens_configurados' not in st.session_state:
    st.session_state['itens_configurados'] = []

# Função para adicionar itens ao resumo
def adicionar_item(descricao, qtde, valor):
    # Substitui " - " por "\n- " para criar quebras de linha na descrição
    descricao_formatada = descricao.replace(" - ", "\n- ")
    total_valor = qtde * valor
    st.session_state.itens_configurados.append({
        "Item": len(st.session_state.itens_configurados) + 1,
        "Qtde": qtde,
        "Descrição": descricao_formatada,
        "Valor Unitário": valor,
        "Valor Total": total_valor
    })

# Função para limpar a tabela de itens configurados
def limpar_tabela():
    st.session_state['itens_configurados'] = []

# Seção 1: Escolha do Tipo de Proposta
st.subheader('Escolha do Tipo de Proposta')
tipo_proposta = st.selectbox(
    'Selecione o tipo de proposta:',
    ['Selecione', 'Subestação Unitária', 'Estação Fotovoltaica']
)

# Configuração básica para Subestação e Estação
if tipo_proposta != 'Selecione':
    st.subheader(f'Configuração da {tipo_proposta}')

    tipo_configuracao = st.selectbox(
        'Selecione a configuração:',
        ['SKID (Transformador Isolado + QBGT)', 'QGBT + Transformador Isolado']
    )

    if tipo_configuracao == 'SKID (Transformador Isolado + QBGT)':
        with st.expander('Configurar SKID'):
            potencia_skid = st.number_input('Potência do Transformador Isolado (kVA):', min_value=0, value=500)
            classe_tensao_skid = st.selectbox('Classe de Tensão do Transformador Isolado (kV):', ['15kV', '24kV', '36kV'])
            perdas_skid = st.text_input('Perdas (%):', value='1,2%')
            fator_k_skid = st.text_input('Fator K:', value='4')
            ip_skid = st.text_input('IP:', value='00')
            corrente_skid = st.number_input('Corrente do QBGT (A):', min_value=0, value=630)
            tensao_skid = st.number_input('Tensão do QBGT (V):', min_value=0, value=800)
            capacidade_curto_skid = st.number_input('Capacidade de Curto-Circuito (kA):', min_value=0, value=10)
            descricao_skid = f"SKID {potencia_skid} kVA\n- SKID ESTRUTURADO EM AÇO PARA {potencia_skid} kVA;\n- TRANSFORMADOR ISOLADO A SECO {potencia_skid} kVA {classe_tensao_skid} kV\n- PERDAS = {perdas_skid}; K = {fator_k_skid}; IP = {ip_skid};\n- QGBT {tensao_skid} V {corrente_skid} A {capacidade_curto_skid} kA"
            quantidade_skid = st.number_input('Quantidade de SKID:', min_value=1, value=1)
            valor_skid = st.number_input('Valor do SKID:', min_value=0.0, format="%.2f")
            if st.button('Adicionar SKID Configurado'):
                adicionar_item(descricao_skid, quantidade_skid, valor_skid)

    elif tipo_configuracao == 'QGBT + Transformador Isolado':
        with st.expander('Configurar QGBT + Transformador Isolado', expanded=False):
            potencia_trafo = st.number_input('Potência (kVA):', min_value=0, value=500)
            classe_tensao_trafo = st.selectbox('Classe de Tensão:', ['15kV', '24kV', '36kV'])
            perdas_trafo = st.text_input('Perdas (%):', value='1,2%')
            fator_k_trafo = st.text_input('Fator K:', value='4')
            ip_trafo = st.text_input('IP:', value='00')
            descricao_trafo = f"TRANSFORMADOR ISOLADO A SECO {potencia_trafo} kVA {classe_tensao_trafo} kV - PERDAS = {perdas_trafo}; K = {fator_k_trafo}; IP = {ip_trafo};"
            quantidade_trafo = st.number_input('Quantidade de Transformadores:', min_value=1, value=1)
            valor_trafo = st.number_input('Valor do Transformador:', min_value=0.0, format="%.2f")
            corrente_qgbt = st.number_input('Corrente (A):', min_value=0, value=630)
            tensao_qgbt = st.number_input('Tensão (V):', min_value=0, value=800)
            capacidade_curto_qgbt = st.number_input('Capacidade de Curto-Circuito (kA):', min_value=0, value=10)
            descricao_qgbt = f"QGBT {tensao_qgbt}V {corrente_qgbt}A {capacidade_curto_qgbt}kA"
            quantidade_qgbt = st.number_input('Quantidade de QBGT:', min_value=1, value=1)
            valor_qgbt = st.number_input('Valor do QBGT:', min_value=0.0, format="%.2f")
            if st.button('Adicionar QGBT + Transformador Configurados'):
                adicionar_item(descricao_trafo, quantidade_trafo, valor_trafo)
                adicionar_item(descricao_qgbt, quantidade_qgbt, valor_qgbt)

        # Lógica da Cabine apenas se "QGBT + Transformador Isolado" for selecionado
        with st.expander('Configurar Cabine'):
            modelo_cabine = st.text_input('Concessionária:')
            quantidade_cabine = st.number_input('Quantidade:', min_value=1, value=1)
            valor_cabine = st.number_input('Valor:', min_value=0.0, format="%.2f")
            if st.button('Adicionar Cabine'):
                descricao_cabine = f"Cabine {modelo_cabine}"
                adicionar_item(descricao_cabine, quantidade_cabine, valor_cabine)

    # Seção para Itens Adicionais apenas se for Estação Fotovoltaica
    if tipo_proposta == 'Estação Fotovoltaica':
        st.subheader('Configuração dos Itens Adicionais')

    # Seção para Inversor
        with st.expander('Configurar Inversor'):
            marca_inversor = st.selectbox('Marca do Inversor', ['Canadian', 'GoodWe', 'Growatt', 'Huawei', 'Sungrow', 'Solis', 'Outras'])
            potencia_inversor = st.selectbox('Potência do Inversor (kW)', [75, 125, 200, 215, 250, 333, 350, 1100, 3125])
            acessorio_inversor = st.selectbox('Acessório do Inversor', ['COM100A', 'COM100E', 'PVS-16MH', 'Logger 3000A', 'Logger 3000B'])
            quantidade_inversor = st.number_input('Quantidade de Inversor:', min_value=1)
            valor_inversor = st.number_input('Valor do Inversor:', min_value=0.0, format="%.2f")

        if st.button('Adicionar Inversor Configurado'):
            descricao_inversor = f"Inversor {marca_inversor} {potencia_inversor} kW com {acessorio_inversor}"
            adicionar_item(descricao_inversor, quantidade_inversor, valor_inversor)

    # Seção para Módulo
        with st.expander('Configurar Módulo'):
            marca_modulo = st.selectbox('Marca do Módulo', ['AE Solar', 'Canadian', 'JA Solar', 'Jinko', 'Longi', 'Risen', 'Trina', 'TW', 'Outras'])
            potencia_modulo = st.selectbox('Potência do Módulo (W)', [450, 455, 540, 545, 550, 555, 560, 565, 570, 575, 600, 650, 655, 660, 665, 680, 690, 695])
            quantidade_modulo = st.number_input('Quantidade de Módulo:', min_value=1)
            valor_modulo = st.number_input('Valor do Módulo:', min_value=0.0, format="%.2f")

        if st.button('Adicionar Módulo Configurado'):
            descricao_modulo = f"Módulo {marca_modulo} {potencia_modulo} W"
            adicionar_item(descricao_modulo, quantidade_modulo, valor_modulo)

    # Seção para Cabine
        with st.expander('Configurar Cabine'):
            fabricante_cabine = st.selectbox('Fabricante da Cabine', ['Blutrafos', 'Gazquez', 'Lucy', 'Progressul', 'Setta', 'N/A'])
            modelo_medicao_cabine = st.selectbox('Modelo de Medição', ['1 Medição', '2 Medições', '3 Medições'])
            quantidade_cabine = st.number_input('Quantidade de Cabine:', min_value=1)
            valor_cabine = st.number_input('Valor da Cabine:', min_value=0.0, format="%.2f")

        if st.button('Adicionar Cabine Configurada'):
            descricao_cabine = f"Cabine {fabricante_cabine} com {modelo_medicao_cabine}"
            adicionar_item(descricao_cabine, quantidade_cabine, valor_cabine)

    # Seção para Estrutura
        with st.expander('Configurar Estrutura'):
            tipo_estrutura = st.selectbox('Tipo de Estrutura', ['Fixa', 'Tracker', 'Flutuante', 'N/A'])
            fixacao_tracker = st.selectbox('Fixação/Tracker', ['Axial', 'Brafer', 'Brametal', 'Dynamo', 'Isoeste', 'Metal Light', 'NTC Somar', 'SSM', 'Valmont', 'Outros', 'N/A'])
            quantidade_estrutura = st.number_input('Quantidade de Estrutura:', min_value=1)
            valor_estrutura = st.number_input('Valor da Estrutura:', min_value=0.0, format="%.2f")

        if st.button('Adicionar Estrutura Configurada'):
            descricao_estrutura = f"Estrutura {tipo_estrutura} com fixação {fixacao_tracker}"
            adicionar_item(descricao_estrutura, quantidade_estrutura, valor_estrutura)

    # Seção para Cabo
        with st.expander('Configurar Cabo'):
            marca_cabo = st.selectbox('Marca do Cabo', ['Cabelauto', 'Conduspar', 'Halukabel', 'New Cabos', 'Nexans', 'Prysmian', 'Reicon', 'N/A'])
            tipo_cabo = st.selectbox('Tipo de Cabo', ['Média Tensão', 'Solar 4mm²', 'Solar 6mm²', 'Solar 8mm²', 'Solar 10mm²'])
            quantidade_cabo = st.number_input('Quantidade de Cabo:', min_value=1)
            valor_cabo = st.number_input('Valor do Cabo:', min_value=0.0, format="%.2f")

        if st.button('Adicionar Cabo Configurado'):
            descricao_cabo = f"Cabo {marca_cabo} {tipo_cabo}"
            adicionar_item(descricao_cabo, quantidade_cabo, valor_cabo)

# Função para remover itens da tabela
    def remover_item(index):
        if 'itens_configurados' in st.session_state:
            st.session_state.itens_configurados.pop(index)

# Exibindo a tabela de itens configurados com opção de remover
st.subheader('Tabela Resumo de Itens Adicionais')

if st.session_state.itens_configurados:
    df = pd.DataFrame(st.session_state.itens_configurados)
    df['Remover'] = df.index
    st.table(df.drop(columns=['Remover']))

    # Adicionando botões de remoção para cada item
    for index, row in df.iterrows():
        if st.button(f'Remover Item {index + 1}'):
            remover_item(index)
else:
    st.write("Nenhum item adicional configurado ainda.")


# Seção Final: Visualização e Resumo
st.subheader('Tabela Resumo de Itens')

# Renderizando a tabela com st.table para melhor quebra de linha
if st.session_state.itens_configurados:
    df = pd.DataFrame(st.session_state.itens_configurados)
    df['Valor Unitário'] = df['Valor Unitário'].apply(lambda x: f"R$ {x:,.2f}".replace(",", "X").replace(".", ",").replace("X", "."))
    df['Valor Total'] = df['Valor Total'].apply(lambda x: f"R$ {x:,.2f}".replace(",", "X").replace(".", ",").replace("X", "."))

    # Calcula o total geral
    total_geral = df['Valor Total'].apply(lambda x: float(x.replace("R$ ", "").replace(".", "").replace(",", "."))).sum()
    
    # Adiciona uma linha de total
    total_df = pd.DataFrame([["", "", "TOTAL", "", f"R$ {total_geral:,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")]], 
                            columns=df.columns)
    df = pd.concat([df, total_df], ignore_index=True)
    
    # Exibindo a tabela usando st.table
    st.table(df)
else:
    st.write("Nenhum item configurado ainda.")

# Botão para limpar a tabela
if st.button('Limpar Tabela'):
    limpar_tabela()

 # Marco 2

# Inicializa a sessão para armazenar eventos de pagamento
if 'eventos_pagamento' not in st.session_state:
    st.session_state['eventos_pagamento'] = {}

# Função para adicionar um evento de pagamento
def adicionar_evento(item_id, percentual, dias, campo_referencia):
    if item_id not in st.session_state['eventos_pagamento']:
        st.session_state['eventos_pagamento'][item_id] = []
    evento = f"{percentual}% - {dias} dias após {campo_referencia}"
    st.session_state['eventos_pagamento'][item_id].append(evento)

# Exibindo a tabela de itens configurados
st.subheader('Condições de Pagamento')
for index, item in enumerate(st.session_state['itens_configurados'], start=1):
    # Filtro para ignorar itens que começam com "Transformador"
    if item['Descrição'].startswith("TRANSFORMADOR"):
        continue
    
    # Título personalizado para QGBT
    if item['Descrição'].split(' ')[0] == "QGBT":
        titulo_item = f"ITEM {index:02d}: QGBT + Transformador Isolado"
    else:
        titulo_item = f"ITEM {index:02d}: {item['Descrição'].split(' ')[0]}"
    
    with st.expander(f"{titulo_item}", expanded=False):
        # Inputs para eventos de pagamento
        percentual = st.number_input(f"Percentual do Evento para {titulo_item}:", min_value=0, max_value=100, value=0, key=f"percentual_{item['Item']}")
        dias = st.number_input(f"Dias para {titulo_item}:", min_value=0, value=0, key=f"dias_{item['Item']}")
        campo_referencia = st.selectbox(f"Campo de Referência para {titulo_item}:", ["Aprovação", "Entrega", "Faturamento", "TAF", "Contraembarque"], key=f"referencia_{item['Item']}")

        # Botão para adicionar evento de pagamento
        if st.button(f"Adicionar Evento para {titulo_item}"):
            adicionar_evento(item['Item'], percentual, dias, campo_referencia)

        # Exibindo eventos adicionados para este item
        if item['Item'] in st.session_state['eventos_pagamento']:
            st.write("**Eventos de Pagamento:**")
            for evento in st.session_state['eventos_pagamento'][item['Item']]:
                st.write(f"- {evento}")






def configurar_impostos(tipo_proposta, icms_subestacao=0):
    if tipo_proposta == "Estação Fotovoltaica":
        st.subheader("Impostos - Estação Fotovoltaica")
        st.write("**NCM**: 8501.72.90 – Sistema Gerador Fotovoltaico")
        st.write("**PIS**: 1,65% inclusos nos preços")
        st.write("**COFINS**: 7,60% inclusos nos preços")
        st.write("**ICMS**: 0,00% inclusos nos preços")
        st.write("**IPI**: 0,00% a incluir nos preços")
    elif tipo_proposta == "Subestação Unitária":
        st.subheader("Impostos - Subestação Unitária")
        icms_subestacao = st.text_input("ICMS (inserir %):", value=str(icms_subestacao))
        st.write("**NCM**: 8537.20.90 – Subestação Unitária")
        st.write("**PIS**: 1,65% inclusos nos preços")
        st.write("**COFINS**: 7,6% inclusos nos preços")
        st.write(f"**ICMS**: {icms_subestacao}% inclusos nos preços")
        st.write("**IPI**: 0,00% a incluir nos preços")

# Aplicação da lógica de impostos baseada no tipo de proposta selecionado
if tipo_proposta != 'Selecione':
    configurar_impostos(tipo_proposta)

st.subheader("Entrega")

tipo_frete = st.radio("Tipo de Frete", options=["CIF", "FOB"])

# Armazena a mensagem correspondente ao tipo de frete escolhido em uma variável
if tipo_frete == "FOB":
    mensagem_frete = "O transporte dos equipamentos comprados será realizado no formato FOB pela empresa Blutrafos, na cidade de Blumenau em Santa Catarina."
elif tipo_frete == "CIF":
    mensagem_frete = "O transporte dos equipamentos comprados será entregue em obra com descarga por conta do cliente no formato CIF pela empresa Blutrafos, na cidade de Blumenau em Santa Catarina."

# Exibe a mensagem correspondente ao tipo de frete escolhido
st.markdown(mensagem_frete)
# Seção: Condição de Entrega

## Seção: Condição de Entrega

if 'tipo_configuracao' in locals() or 'tipo_configuracao' in globals():
    # Exibindo a tabela de itens configurados
    st.subheader('Condições de Entrega')

    # Prazos pré-estabelecidos
    prazo_desenhos_aprovacao = st.number_input('Desenhos para aprovação (dias):', min_value=0, format="%d")
    prazo_aprovacao_cliente = st.number_input('Prazo para aprovação dos desenhos pelo cliente (dias):', min_value=0, format="%d")

    # Inicializa a lista para os prazos dos itens
    prazos_itens = []

    for index, item in enumerate(st.session_state['itens_configurados'], start=1):
        # Filtro para ignorar itens que começam com "TRANSFORMADOR"
        if item['Descrição'].startswith("TRANSFORMADOR"):
            continue

        # Título personalizado para QGBT
        if item['Descrição'].split(' ')[0] == "QGBT":
            titulo_item = f"ITEM {index:02d}: QGBT + Transformador Isolado"
        else:
            titulo_item = f"ITEM {index:02d}: {item['Descrição'].split(' ')[0]}"

        with st.expander(f"{titulo_item}", expanded=False):
            # Inputs para prazos de entrega
            prazo_entrega = st.text_input(f"Prazo de entrega para {titulo_item} (dias):", key=f"prazo_{item['Item']}")
            prazos_itens.append((titulo_item, prazo_entrega))

    # Seção Final: Resumo em Texto
    resumo_texto = "### Resumo Entrega\n"
    resumo_texto += f"- Desenhos para aprovação: {prazo_desenhos_aprovacao} dias\n"
    resumo_texto += f"- Prazo para aprovação dos desenhos pelo cliente: {prazo_aprovacao_cliente} dias\n\n"

    for titulo, prazo in prazos_itens:
        resumo_texto += f"- Prazo de entrega para {titulo}: {prazo} dias\n"

    # Exibe o Resumo
    st.markdown(resumo_texto)

else:
    st.warning("Por favor, selecione uma configuração antes de definir os prazos de entrega.")

    # Resumo de Comentários
st.subheader('Resumo de Comentários')

# Inicializa a sessão para armazenar os comentários se ainda não estiver presente
if 'comentarios' not in st.session_state:
    st.session_state['comentarios'] = []

# Função para adicionar um comentário
def adicionar_comentario(comentario):
    st.session_state['comentarios'].append(comentario)

# Caixa de texto para adicionar um novo comentário
novo_comentario = st.text_area('Adicione um comentário:')

# Botão para adicionar o comentário
if st.button('Adicionar Comentário'):
    if novo_comentario:
        adicionar_comentario(novo_comentario)
        st.success('Comentário adicionado com sucesso!')
    else:
        st.warning('O comentário não pode ser vazio.')

# Exibe o resumo de comentários
if st.session_state['comentarios']:
    resumo_comentarios = "### Comentários Adicionados\n"
    for i, comentario in enumerate(st.session_state['comentarios'], start=1):
        resumo_comentarios += f"**Comentário {i}:** {comentario}\n"
    
    st.markdown(resumo_comentarios)
else:
    st.write("Nenhum comentário adicionado ainda.")

garantia= st.text('Adicione a garantia:')

validade= st.text('Adicione a validade:')

import pythoncom
from win32com.client import Dispatch
import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# Seleção do responsável
responsavel = st.selectbox('Responsável:', ['Gabrielle Visintainer', 'Bernardo Seelig Orige', 'Paulo Fernando Knoch'])

# Seleção do gerente
gerente = st.selectbox('Gerente:', ['Tatiane Bendotti', 'Mauricio Saldanha'])

# Dicionário com as informações de contato de cada responsável e gerente
contatos_divisao_solar = {
    "Gabrielle Visintainer": {
        "Nome": "Gabrielle Visintainer",
        "Cargo": "COMERCIAL – Divisão Solar",
        "Telefone": "(47) 3036-3004",
        "Celular": "(47) 99207-2067",
        "Emails": ["gabrielle.pinto@blutrafos.com.br", "vendas@blutrafos.com.br"]
    },
    "Bernardo Seelig Orige": {
        "Nome": "Bernardo Seelig Orige",
        "Cargo": "COMERCIAL – Divisão Solar",
        "Celular": "(47) 99782-0204",
        "Emails": ["bernardo.orige@blutrafos.com.br", "vendas@blutrafos.com.br"]
    },
    "Paulo Fernando Knoch": {
        "Nome": "Paulo Fernando Knoch",
        "Cargo": "COMERCIAL – Divisão Solar",
        "Celular": "(47) 99207-2067",
        "Emails": ["paulo.knoch@blutrafos.com.br", "vendas@blutrafos.com.br"]
    },
    "Tatiane Bendotti": {
        "Nome": "Tatiane Bendotti",
        "Cargo": "GERENTE DE NEGÓCIOS – Divisão Solar",
        "Telefone": "(47) 3036-3000",
        "Celular": "(47) 99133-4539",
        "Emails": ["tatiane@blutrafos.com.br", "vendas@blutrafos.com.br"]
    },
    "Mauricio Saldanha": {
        "Nome": "Mauricio Saldanha",
        "Cargo": "GERENTE DE NEGÓCIOS – Divisão Solar SP",
        "Celular": "(47) 99735-0290",
        "Emails": ["mauricio@blutrafos.com.br", "vendas@blutrafos.com.br"]
    }
}

# Função para aplicar formatação a um parágrafo
def apply_paragraph_formatting(paragraph, alignment='left', space_before=Pt(0), space_after=Pt(0)):
    paragraph_format = paragraph.paragraph_format
    paragraph_format.space_before = space_before
    paragraph_format.space_after = space_after
    paragraph_format.line_spacing = 1  # Espaçamento simples
    
    if alignment == 'left':
        paragraph.alignment = 0  # Alinhamento à esquerda
    elif alignment == 'center':
        paragraph.alignment = 1  # Alinhamento centralizado
    elif alignment == 'right':
        paragraph.alignment = 2  # Alinhamento à direita

# Função para aplicar sombreamento a uma célula
def set_cell_shading(cell, color):
    tc_pr = cell._element.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color)
    tc_pr.append(shd)

# Função para criar a tabela personalizada
def create_custom_table(doc, itens_configurados, tipo_proposta):
    num_linhas = len(itens_configurados) + 2
    table = doc.add_table(rows=num_linhas, cols=4)

    widths = [Inches(0.5), Inches(0.5), Inches(4.5), Inches(1.5)]
    for idx, width in enumerate(widths):
        for cell in table.columns[idx].cells:
            cell.width = width

    header_row = table.rows[0]
    header_data = ["Item", "Qtde", f"Escopo de Fornecimento: {tipo_proposta}", "Valor Total"]
    
    for idx, cell in enumerate(header_row.cells):
        cell.text = header_data[idx]
        paragraph = cell.paragraphs[0]
        run = paragraph.runs[0]
        run.font.name = 'Calibri Light'
        run.font.size = Pt(11)
        run.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
        apply_paragraph_formatting(paragraph, alignment='center', space_before=Pt(0.25), space_after=Pt(0.25))
        
        tc = cell._element
        tcPr = tc.get_or_add_tcPr()
        tcBorders = OxmlElement('w:tcBorders')
        for border_name in ['top', 'left', 'bottom', 'right']:
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), 'double')
            border.set(qn('w:sz'), '4')
            border.set(qn('w:space'), '0')
            border.set(qn('w:color'), '000000')
            tcBorders.append(border)
        tcPr.append(tcBorders)
        
        set_cell_shading(cell, '00543C')

    for idx, item in enumerate(itens_configurados, start=1):
        row = table.rows[idx]
        row.cells[0].text = str(item["Item"])
        row.cells[1].text = str(item["Qtde"])
        row.cells[2].text = item["Descrição"]
        row.cells[3].text = f"R$ {item['Valor Total']:,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")

        for cell in row.cells:
            paragraph = cell.paragraphs[0]
            run = paragraph.runs[0]
            run.font.name = 'Calibri Light'
            run.font.size = Pt(11)

            if cell == row.cells[2]:
                apply_paragraph_formatting(paragraph, alignment='left')
            else:
                apply_paragraph_formatting(paragraph, alignment='center')

            tc = cell._element
            tcPr = tc.get_or_add_tcPr()
            tcBorders = OxmlElement('w:tcBorders')
            for border_name in ['top', 'left', 'bottom', 'right']:
                border = OxmlElement(f'w:{border_name}')
                border.set(qn('w:val'), 'double')
                border.set(qn('w:sz'), '4')
                border.set(qn('w:space'), '0')
                border.set(qn('w:color'), '000000')
                tcBorders.append(border)
            tcPr.append(tcBorders)

    last_row = table.rows[-1]
    last_row.cells[0].merge(last_row.cells[1])
    last_row.cells[0].merge(last_row.cells[2])
    last_row.cells[0].merge(last_row.cells[3])
    
    total_geral = sum(item['Valor Total'] for item in itens_configurados)
    last_row.cells[0].text = f"TOTAL (s/IPI): R$ {total_geral:,.2f}".replace(",", "X").replace(".", ",").replace("X", ",")
    last_row.cells[0].paragraphs[0].runs[0].font.name = 'Calibri Light'
    last_row.cells[0].paragraphs[0].runs[0].font.size = Pt(11)
    last_row.cells[0].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
    apply_paragraph_formatting(last_row.cells[0].paragraphs[0], alignment='right', space_before=Pt(0.25), space_after=Pt(0.25))
    
    set_cell_shading(last_row.cells[0], '00543C')

    return table

# Função para inserir a tabela abaixo de um parágrafo específico
def insert_table_below_paragraph(docx_path, paragraph_text, itens_configurados, tipo_proposta, output_path):
    doc = Document(docx_path)

    for i, paragraph in enumerate(doc.paragraphs):
        if paragraph.text.strip() == paragraph_text:
            new_paragraph = doc.paragraphs[i+1].insert_paragraph_before("")

            table = create_custom_table(doc, itens_configurados, tipo_proposta)

            doc.paragraphs[i+1]._element.addnext(table._element)
            break

    doc.save(output_path)

# Função para substituir o placeholder de impostos por texto formatado
def replace_varImposto(replacements, tipo_proposta):
    if tipo_proposta == "Subestação Unitária":
        formatted_title = "NCM: 8537.20.90 – Subestação Unitária"
        formatted_text = [
            "✔ PIS: 1,65% inclusos nos preços;",
            "✔ COFINS: 7,6% inclusos nos preços;",
            "✔ ICMS: 12,00% inclusos nos preços;",
            "✔ IPI: 0,00% a incluir nos preços."
        ]
    elif tipo_proposta == "Estação Fotovoltaica":
        formatted_title = "NCM: 8501.72.90 – Sistema Gerador Fotovoltaico"
        formatted_text = [
            "✔ PIS: 1,65% inclusos nos preços;",
            "✔ COFINS: 7,60% inclusos nos preços;",
            "✔ ICMS: 0,00% inclusos nos preços;",
            "✔ IPI: 0,00% a incluir nos preços."
        ]
    
    replacements['_varImposto2'] = (formatted_title, formatted_text)

# Função para substituir texto nos parágrafos e cabeçalhos
def replace_text_in_paragraphs_and_headers(docx_path, replacements, output_path, tipo_proposta):
    doc = Document(docx_path)

    for paragraph in doc.paragraphs:
        for old_text, new_text in replacements.items():
            if old_text in paragraph.text:
                if old_text == '_varImposto2':
                    formatted_title, formatted_text = new_text
                    
                    # Substitui o texto do parágrafo com o título formatado
                    paragraph.text = formatted_title
                    paragraph.style = doc.styles["Heading 2"]
                    
                    # Adiciona as linhas formatadas abaixo do título
                    for line in formatted_text:
                        new_paragraph = paragraph.insert_paragraph_before(line)
                        new_paragraph.style = doc.styles["List Paragraph"]
                        new_paragraph.paragraph_format.space_after = Pt(0)
                        new_paragraph.paragraph_format.line_spacing = 1
                    paragraph = new_paragraph
                else:
                    paragraph.text = paragraph.text.replace(old_text, new_text)
    
    for section in doc.sections:
        header = section.header
        for paragraph in header.paragraphs:
            for old_text, new_text in replacements.items():
                if old_text in paragraph.text:
                    paragraph.text = paragraph.text.replace(old_text, new_text)

    doc.save(output_path)

# Função para substituir texto nas caixas de texto (Shapes)
def replace_text_in_textboxes(docx_path, replacements, output_path):
    pythoncom.CoInitialize()
    word = Dispatch("Word.Application")
    word.Visible = False

    doc = word.Documents.Open(docx_path)

    for shape in doc.Shapes:
        if shape.TextFrame.HasText:
            text_range = shape.TextFrame.TextRange
            for old_text, new_text in replacements.items():
                if old_text in text_range.Text:
                    text_range.Text = text_range.Text.replace(old_text, new_text)

    doc.SaveAs(output_path)
    doc.Close(False)
    word.Quit()

# Função para remover as informações dos contatos não selecionados
def remove_unselected_contacts(doc, selected_responsavel, selected_gerente):
    paragraphs_to_keep = []
    paragraphs_to_remove = []
    
    for paragraph in doc.paragraphs:
        if selected_responsavel not in paragraph.text and selected_gerente not in paragraph.text:
            if any(contact in paragraph.text for contact in contatos_divisao_solar.keys()):
                paragraphs_to_remove.append(paragraph)
            else:
                paragraphs_to_keep.append(paragraph)
        else:
            paragraphs_to_keep.append(paragraph)
    
    # Remove os parágrafos indesejados
    for paragraph in paragraphs_to_remove:
        paragraph.clear()

# Integração no botão de Streamlit
if st.button('Gerar Documento com Tabela'):
    docx_path = 'Template_Proposta_Comercial.docx'
    output_path = r'C:\Users\Erick\Desktop\Automacao_Proposta_Solar\documento_modificado.docx'
    itens_configurados = st.session_state.itens_configurados  # Dados dinâmicos do Streamlit
    # tipo_proposta já é definido anteriormente no seu código
    # Usa o tipo_proposta diretamente para as substituições e formatações

    # Primeiro, insere a tabela abaixo do parágrafo desejado
    insert_table_below_paragraph(docx_path, "Quadro de Preços", itens_configurados, tipo_proposta, output_path)
    
    # Salva o documento antes de realizar a substituição
    docx_path_after_table = output_path
    
    # Substituições a serem feitas no documento
    replacements = {
        '_varCliente': str(cliente),
        '_varObra': str(obra),
        '_varOR': str(numero_or),
        '_varRev': str(revisao),
        '_varDia': str(dia),
        '_varMes': str(mes),
        '_varAno': str(ano),
        'varDolar': str(cotacao_dolar),
        '_varGarantia': str(garantia),
        '_varValidade': str(validade),
        '_varTransporte': str(mensagem_frete)
    }
    
    # Adiciona a substituição do _varImposto2
    replace_varImposto(replacements, tipo_proposta)
    
    # Substitui o texto nas caixas de texto
    replace_text_in_textboxes(docx_path_after_table, replacements, output_path)
    
    # Substitui o texto nos parágrafos e cabeçalhos
    replace_text_in_paragraphs_and_headers(output_path, replacements, output_path, tipo_proposta)

    # Abre o documento Word
    doc = Document(output_path)

    # Remove os contatos que não foram selecionados
    remove_unselected_contacts(doc, responsavel, gerente)
    
    # Salva o documento após as alterações
    doc.save(output_path)

    # Abre o documento automaticamente
    os.startfile(output_path)

    st.success("Tabela inserida, texto substituído, e documento salvo como 'documento_modificado.docx'")
