import os
import streamlit as st
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import pythoncom
from win32com.client import Dispatch
from datetime import datetime

def configurar_informacoes():
    # Configuração dos dados iniciais
    st.subheader('Dados Iniciais')

    # Capturando a data de hoje como valor padrão
    data_hoje = datetime.today()
    data_selecionada = st.date_input('Data da Proposta:', value=data_hoje)

    # Armazenando os dados no session_state
    st.session_state['dados_iniciais'] = {
        'cliente': st.text_input('Cliente:'),
        'obra': st.text_input('Obra:'),
        'numero_or': st.text_input('Número da OR:'),
        'revisao': st.text_input('Revisão:'),
        'dia': data_selecionada.strftime('%d'),  # Dia selecionado
        'mes': data_selecionada.strftime('%B'),  # Nome completo do mês selecionado
        'ano': data_selecionada.strftime('%Y'),  # Ano selecionado
        'cotacao_dolar': st.number_input('Cotação do Dólar:', min_value=0.0, value=5.4, format="%.2f")
    }

    # Configuração do tipo de frete
    st.subheader('Configuração do Frete')
    tipo_frete = st.radio("Tipo de Frete", options=["CIF", "FOB"])
    st.session_state['tipo_frete'] = tipo_frete

    # Seleção do responsável e gerente
    st.subheader('Responsável e Gerente')
    responsavel_selecionado = st.selectbox('Responsável:', ['Gabrielle Visintainer', 'Bernardo Seelig Orige', 'Paulo Fernando Knoch'])
    gerente_selecionado = st.selectbox('Gerente:', ['Tatiane Bendotti', 'Mauricio Saldanha'])
    st.session_state['responsavel'] = responsavel_selecionado
    st.session_state['gerente'] = gerente_selecionado

    # Configuração da garantia e validade
    st.subheader('Garantia e Validade')
    st.session_state['garantia'] = st.text_input('Garantia:')
    st.session_state['validade'] = st.text_input('Validade:')

def gerar_documento():
    docx_path = 'Template_Proposta_Comercial.docx'
    output_path = 'documento_modificado.docx'
    itens_configurados = st.session_state.itens_configurados
    tipo_proposta = st.session_state['tipo_proposta']

    # Dados de contatos para substituição
    contatos_divisao_solar = {
        "Gabrielle Visintainer": {
            "Nome": "Gabrielle Visintainer",
            "Cargo": "COMERCIAL – Divisão Solar",
            "Telefone": "(47) 3036-3004",
            "Celular": "(47) 99207-2067",
            "Emails": ["gabrielle.pinto@blutrafos.com.br", "vendas@blutrafos.com.br"]
        },
        "Bernardo Seelig Orige": {
            "Nome": "Bernardo Seelig Orige",
            "Cargo": "COMERCIAL – Divisão Solar",
            "Celular": "(47) 99782-0204",
            "Emails": ["bernardo.orige@blutrafos.com.br", "vendas@blutrafos.com.br"]
        },
        "Paulo Fernando Knoch": {
            "Nome": "Paulo Fernando Knoch",
            "Cargo": "COMERCIAL – Divisão Solar",
            "Celular": "(47) 99207-2067",
            "Emails": ["paulo.knoch@blutrafos.com.br", "vendas@blutrafos.com.br"]
        },
        "Tatiane Bendotti": {
            "Nome": "Tatiane Bendotti",
            "Cargo": "GERENTE DE NEGÓCIOS – Divisão Solar",
            "Telefone": "(47) 3036-3000",
            "Celular": "(47) 99133-4539",
            "Emails": ["tatiane@blutrafos.com.br", "vendas@blutrafos.com.br"]
        },
        "Mauricio Saldanha": {
            "Nome": "Mauricio Saldanha",
            "Cargo": "GERENTE DE NEGÓCIOS – Divisão Solar SP",
            "Celular": "(47) 99735-0290",
            "Emails": ["mauricio@blutrafos.com.br", "vendas@blutrafos.com.br"]
        }
    }

    # Replacements para os placeholders [varresponsavel1] e [vargerente]
    responsavel = contatos_divisao_solar[st.session_state['responsavel']]
    gerente = contatos_divisao_solar[st.session_state['gerente']]

    replacements = {
        '_varCliente': st.session_state['dados_iniciais']['cliente'],
        '_varObra': st.session_state['dados_iniciais']['obra'],
        '_varOR': st.session_state['dados_iniciais']['numero_or'],
        '_varRev': st.session_state['dados_iniciais']['revisao'],
        '_varDia': st.session_state['dados_iniciais']['dia'],
        '_varMes': st.session_state['dados_iniciais']['mes'],
        '_varAno': st.session_state['dados_iniciais']['ano'],
        'varDolar': st.session_state['dados_iniciais']['cotacao_dolar'],
        '_varGarantia': st.session_state['garantia'],
        '_varValidade': st.session_state['validade'],
        '_varTransporte': f"Frete: {st.session_state['tipo_frete']}",
        '[varresponsavel1]': f"{responsavel['Nome']}\n{responsavel['Cargo']}\nTelefone: {responsavel.get('Telefone', 'N/A')}\nCelular: {responsavel['Celular']}\nEmails: {', '.join(responsavel['Emails'])}",
        '[vargerente]': f"{gerente['Nome']}\n{gerente['Cargo']}\nTelefone: {gerente.get('Telefone', 'N/A')}\nCelular: {gerente['Celular']}\nEmails: {', '.join(gerente['Emails'])}"
    }
    
    # Substituição da variável _varImposto2
    replace_varImposto(replacements, tipo_proposta)
    
    # Substitui o texto nas caixas de texto
    replace_text_in_textboxes(docx_path, replacements, output_path)
    
    # Substitui o texto nos parágrafos e cabeçalhos
    replace_text_in_paragraphs_and_headers(output_path, replacements, output_path, tipo_proposta)

    # Abre o documento Word
    doc = Document(output_path)

    # Salva o documento após as alterações
    doc.save(output_path)

    # Abre o documento automaticamente
    os.startfile(output_path)

    st.success("Tabela inserida, texto substituído, e documento salvo como 'documento_modificado.docx'")

def insert_table_below_paragraph(docx_path, paragraph_text, itens_configurados, tipo_proposta, output_path):
    doc = Document(docx_path)
    for i, paragraph in enumerate(doc.paragraphs):
        if paragraph.text.strip() == paragraph_text:
            new_paragraph = doc.paragraphs[i+1].insert_paragraph_before("")
            table = create_custom_table(doc, itens_configurados, tipo_proposta)
            doc.paragraphs[i+1]._element.addnext(table._element)
            break
    doc.save(output_path)

def create_custom_table(doc, itens_configurados, tipo_proposta):
    num_linhas = len(itens_configurados) + 2
    table = doc.add_table(rows=num_linhas, cols=4)

    # Definir larguras das colunas
    widths = [Inches(0.5), Inches(0.5), Inches(4.5), Inches(1.5)]
    for idx, width in enumerate(widths):
        for cell in table.columns[idx].cells:
            cell.width = width

    # Criar cabeçalho da tabela
    header_row = table.rows[0]
    header_data = ["Item", "Qtde", f"Escopo de Fornecimento: {tipo_proposta}", "Valor Total"]
    
    for idx, cell in enumerate(header_row.cells):
        cell.text = header_data[idx]
        paragraph = cell.paragraphs[0]
        run = paragraph.runs[0]
        run.font.name = 'Calibri Light'
        run.font.size = Pt(11)
        run.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
        apply_paragraph_formatting(paragraph, alignment='center', space_before=Pt(0.25), space_after=Pt(0.25))
        
        tc = cell._element
        tcPr = tc.get_or_add_tcPr()
        tcBorders = OxmlElement('w:tcBorders')
        for border_name in ['top', 'left', 'bottom', 'right']:
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), 'double')
            border.set(qn('w:sz'), '4')
            border.set(qn('w:space'), '0')
            border.set(qn('w:color'), '000000')
            tcBorders.append(border)
        tcPr.append(tcBorders)
        
        set_cell_shading(cell, '00543C')

    # Preencher as linhas da tabela com itens configurados
    for idx, item in enumerate(itens_configurados, start=1):
        row = table.rows[idx]
        row.cells[0].text = str(item["Item"])
        row.cells[1].text = str(item["Qtde"])
        row.cells[2].text = item["Descrição"]
        row.cells[3].text = f"R$ {item['Valor Total']:,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")

        for cell in row.cells:
            paragraph = cell.paragraphs[0]
            run = paragraph.runs[0]
            run.font.name = 'Calibri Light'
            run.font.size = Pt(11)

            if cell == row.cells[2]:
                apply_paragraph_formatting(paragraph, alignment='left')
            else:
                apply_paragraph_formatting(paragraph, alignment='center')

            tc = cell._element
            tcPr = tc.get_or_add_tcPr()
            tcBorders = OxmlElement('w:tcBorders')
            for border_name in ['top', 'left', 'bottom', 'right']:
                border = OxmlElement(f'w:{border_name}')
                border.set(qn('w:val'), 'double')
                border.set(qn('w:sz'), '4')
                border.set(qn('w:space'), '0')
                border.set(qn('w:color'), '000000')
                tcBorders.append(border)
            tcPr.append(tcBorders)

    # Adicionar linha de total geral
    last_row = table.rows[-1]
    last_row.cells[0].merge(last_row.cells[1])
    last_row.cells[0].merge(last_row.cells[2])
    last_row.cells[0].merge(last_row.cells[3])
    
    total_geral = sum(item['Valor Total'] for item in itens_configurados)
    last_row.cells[0].text = f"TOTAL (s/IPI): R$ {total_geral:,.2f}".replace(",", "X").replace(".", ",").replace("X", ",")
    last_row.cells[0].paragraphs[0].runs[0].font.name = 'Calibri Light'
    last_row.cells[0].paragraphs[0].runs[0].font.size = Pt(11)
    last_row.cells[0].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
    apply_paragraph_formatting(last_row.cells[0].paragraphs[0], alignment='right', space_before=Pt(0.25), space_after=Pt(0.25))
    
    set_cell_shading(last_row.cells[0], '00543C')

    return table

def replace_varImposto(replacements, tipo_proposta):
    if tipo_proposta == "Subestação Unitária":
        formatted_title = "NCM: 8537.20.90 – Subestação Unitária"
        formatted_text = [
            "✔ PIS: 1,65% inclusos nos preços;",
            "✔ COFINS: 7,6% inclusos nos preços;",
            "✔ ICMS: 12,00% inclusos nos preços;",
            "✔ IPI: 0,00% a incluir nos preços."
        ]
    elif tipo_proposta == "Estação Fotovoltaica":
        formatted_title = "NCM: 8501.72.90 – Sistema Gerador Fotovoltaico"
        formatted_text = [
            "✔ PIS: 1,65% inclusos nos preços;",
            "✔ COFINS: 7,60% inclusos nos preços;",
            "✔ ICMS: 0,00% inclusos nos preços;",
            "✔ IPI: 0,00% a incluir nos preços."
        ]
    replacements['_varImposto2'] = (formatted_title, formatted_text)

def replace_text_in_textboxes(docx_path, replacements, output_path):
    pythoncom.CoInitialize()
    word = Dispatch("Word.Application")
    word.Visible = False
    doc = word.Documents.Open(docx_path)
    for shape in doc.Shapes:
        if shape.TextFrame.HasText:
            text_range = shape.TextFrame.TextRange
            for old_text, new_text in replacements.items():
                if old_text in text_range.Text:
                    text_range.Text = text_range.Text.replace(old_text, new_text)
    doc.SaveAs(output_path)
    doc.Close(False)
    word.Quit()

def replace_text_in_paragraphs_and_headers(docx_path, replacements, output_path, tipo_proposta):
    doc = Document(docx_path)
    for paragraph in doc.paragraphs:
        for old_text, new_text in replacements.items():
            if old_text in paragraph.text:
                if old_text == '_varImposto2':
                    formatted_title, formatted_text = new_text
                    paragraph.text = formatted_title
                    paragraph.style = doc.styles["Heading 2"]
                    for line in formatted_text:
                        new_paragraph = paragraph.insert_paragraph_before(line)
                        new_paragraph.style = doc.styles["List Paragraph"]
                        new_paragraph.paragraph_format.space_after = Pt(0)
                        new_paragraph.paragraph_format.line_spacing = 1
                    paragraph = new_paragraph
                else:
                    paragraph.text = paragraph.text.replace(old_text, new_text)
    for section in doc.sections:
        header = section.header
        for paragraph in header.paragraphs:
            for old_text, new_text in replacements.items():
                if old_text in paragraph.text:
                    paragraph.text = paragraph.text.replace(old_text, new_text)
    doc.save(output_path)

def apply_paragraph_formatting(paragraph, alignment='left', space_before=Pt(0), space_after=Pt(0)):
    paragraph_format = paragraph.paragraph_format
    paragraph_format.space_before = space_before
    paragraph_format.space_after = space_after
    paragraph_format.line_spacing = 1  # Espaçamento simples
    
    if alignment == 'left':
        paragraph.alignment = 0  # Alinhamento à esquerda
    elif alignment == 'center':
        paragraph.alignment = 1  # Alinhamento centralizado
    elif alignment == 'right':
        paragraph.alignment = 2  # Alinhamento à direita

def set_cell_shading(cell, color):
    tc_pr = cell._element.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color)
    tc_pr.append(shd)

# Chamada para configurar as informações na interface
configurar_informacoes()

# Botão para gerar o documento
if st.button('Gerar Documento'):
    gerar_documento()
